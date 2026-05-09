from datetime import date
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(34, 34, 34)

    doc.core_properties.title = "Propuesta KUMAINA - Ministerio TIC"
    doc.core_properties.subject = "Propuesta institucional para preservación de la lengua palenquera"
    doc.core_properties.author = "Codex"
    doc.core_properties.comments = "Documento preparado para articulación institucional"


def configure_header_footer(doc):
    section = doc.sections[0]
    header = section.header
    header_p = header.paragraphs[0]
    header_p.text = "KUMAINA | Ministerio TIC"
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header_p.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(111, 111, 111)

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run("Documento de propuesta institucional")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(111, 111, 111)
    for run in footer_p.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(111, 111, 111)


def add_cover(doc):
    logo_path = Path(__file__).resolve().parents[1] / "assets" / "images" / "brand" / "logo horizontal.png"
    if logo_path.exists():
        logo = doc.add_paragraph()
        logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo.add_run().add_picture(str(logo_path), width=Inches(2.8))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("KUMAINA")
    run.bold = True
    run.font.size = Pt(32)
    run.font.color.rgb = RGBColor(13, 67, 52)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Lengua Palenquera")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(120, 90, 24)

    slogan = doc.add_paragraph()
    slogan.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = slogan.add_run("Habla nuestras raíces.")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(198, 141, 22)

    lead = doc.add_paragraph()
    lead.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lead_run = lead.add_run(
        "Una propuesta hecha con y para jóvenes, para aprender, vivir y preservar la lengua palenquera "
        "desde el celular, con cultura, juego, diccionario digital y traductor comunitario."
    )
    lead_run.font.size = Pt(11)
    lead_run.italic = True
    lead_run.font.color.rgb = RGBColor(83, 83, 83)

    doc.add_paragraph("")

    facts = doc.add_table(rows=2, cols=2)
    facts.style = "Table Grid"
    facts.autofit = True
    fact_rows = [
        ("Dirigido a", "Jóvenes, comunidades y aliados educativos"),
        ("Enfoque", "Preservación cultural + aprendizaje móvil + lexicografía digital"),
        ("Duración", "6 meses para MVP y piloto"),
        ("Estado", "Propuesta para articulación institucional"),
    ]
    idx = 0
    for r in range(2):
        for c in range(2):
            label, value = fact_rows[idx]
            cell = facts.rows[r].cells[c]
            cell.text = f"{label}\n{value}"
            set_cell_shading(cell, "F4E7C5" if c == 0 else "FFFDF6")
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for i, run in enumerate(p.runs):
                    run.font.size = Pt(10)
                    if i == 0:
                        run.bold = True
                        run.font.color.rgb = RGBColor(13, 67, 52)
                    else:
                        run.font.color.rgb = RGBColor(48, 48, 48)
            idx += 1

    doc.add_paragraph("")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        f"Documento de presentación institucional | Ministerio TIC | {date.today().strftime('%d/%m/%Y')}"
    )
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_page_break()


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(13, 67, 52)
    return h


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    return p


def add_soft_callout(doc, text):
    box = doc.add_table(rows=1, cols=1)
    box.style = "Table Grid"
    cell = box.rows[0].cells[0]
    cell.text = text
    set_cell_shading(cell, "FFF7E4")
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            run.italic = True
            run.font.color.rgb = RGBColor(106, 78, 20)


def _load_font(size, bold=False, italic=False):
    candidates = []
    if bold and italic:
        candidates = [
            "C:/Windows/Fonts/segoeuii.ttf",
            "C:/Windows/Fonts/arialbi.ttf",
            "C:/Windows/Fonts/calibriz.ttf",
        ]
    elif bold:
        candidates = [
            "C:/Windows/Fonts/segoeuib.ttf",
            "C:/Windows/Fonts/arialbd.ttf",
            "C:/Windows/Fonts/calibrib.ttf",
        ]
    elif italic:
        candidates = [
            "C:/Windows/Fonts/segoeuii.ttf",
            "C:/Windows/Fonts/ariali.ttf",
            "C:/Windows/Fonts/calibrii.ttf",
        ]
    else:
        candidates = [
            "C:/Windows/Fonts/segoeui.ttf",
            "C:/Windows/Fonts/arial.ttf",
            "C:/Windows/Fonts/calibri.ttf",
        ]

    for candidate in candidates:
        font_path = Path(candidate)
        if font_path.exists():
            return ImageFont.truetype(str(font_path), size=size)
    return ImageFont.load_default()


def _wrap_text(draw, text, font, max_width):
    lines = []
    for raw_line in text.splitlines() or [""]:
        words = raw_line.split()
        if not words:
            lines.append("")
            continue
        current = words[0]
        for word in words[1:]:
            trial = f"{current} {word}"
            if draw.textbbox((0, 0), trial, font=font)[2] <= max_width:
                current = trial
            else:
                lines.append(current)
                current = word
        lines.append(current)
    return lines


def _draw_wrapped_text(draw, xy, text, font, fill, max_width, line_spacing=8):
    x, y = xy
    for line in _wrap_text(draw, text, font, max_width):
        draw.text((x, y), line, font=font, fill=fill)
        bbox = draw.textbbox((x, y), line, font=font)
        y = bbox[3] + line_spacing
    return y


def _draw_progress_bar(draw, x, y, width, value, fill, track=(220, 210, 190, 255), height=18):
    radius = height // 2
    draw.rounded_rectangle((x, y, x + width, y + height), radius=radius, fill=track)
    filled = max(8, int(width * value))
    draw.rounded_rectangle((x, y, x + filled, y + height), radius=radius, fill=fill)


def _draw_card(draw, box, title, subtitle, bullets, progress, accent, tag, highlight=None):
    x1, y1, x2, y2 = box
    shadow = (x1 + 12, y1 + 14, x2 + 12, y2 + 14)
    draw.rounded_rectangle(shadow, radius=34, fill=(0, 0, 0, 60))
    draw.rounded_rectangle(box, radius=34, fill=(246, 236, 216, 255), outline=accent, width=4)

    tag_font = _load_font(18, bold=True)
    title_font = _load_font(32, bold=True)
    subtitle_font = _load_font(20)
    body_font = _load_font(19)
    small_bold = _load_font(18, bold=True)

    draw.rounded_rectangle((x1 + 24, y1 + 22, x1 + 24 + 165, y1 + 22 + 38), radius=18, fill=accent)
    draw.text((x1 + 36, y1 + 28), tag, font=tag_font, fill=(255, 255, 255, 255))
    draw.text((x1 + 24, y1 + 74), title, font=title_font, fill=(13, 67, 52, 255))
    draw.text((x1 + 24, y1 + 118), subtitle, font=subtitle_font, fill=(101, 88, 60, 255))

    text_y = y1 + 160
    if highlight:
        draw.rounded_rectangle(
            (x1 + 24, text_y, x2 - 24, text_y + 94),
            radius=20,
            fill=(231, 246, 239, 255),
            outline=(176, 130, 34, 255),
            width=2,
        )
        draw.text((x1 + 40, text_y + 16), highlight[0], font=_load_font(28, bold=True), fill=(13, 67, 52, 255))
        draw.text((x1 + 40, text_y + 54), highlight[1], font=_load_font(16), fill=(105, 90, 65, 255))
        text_y += 112

    for bullet in bullets:
        draw.ellipse((x1 + 24, text_y + 8, x1 + 36, text_y + 20), fill=accent)
        text_y = _draw_wrapped_text(
            draw,
            (x1 + 48, text_y),
            bullet,
            body_font,
            (40, 40, 40, 255),
            max_width=(x2 - x1) - 88,
            line_spacing=6,
        )
        text_y += 6

    _draw_progress_bar(draw, x1 + 24, y2 - 52, (x2 - x1) - 48, progress, accent)
    pct = f"{int(progress * 100)}%"
    draw.text((x2 - 72, y2 - 79), pct, font=small_bold, fill=(13, 67, 52, 255))


def build_boceto_avance_image(output_path):
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    width, height = 1600, 900
    base = Image.new("RGBA", (width, height), (8, 32, 24, 255))
    overlay = Image.new("RGBA", (width, height), (0, 0, 0, 0))
    od = ImageDraw.Draw(overlay)

    od.ellipse((-260, -180, 760, 760), fill=(24, 88, 56, 110))
    od.ellipse((940, -220, 1840, 560), fill=(187, 141, 38, 80))
    od.ellipse((980, 430, 1840, 1220), fill=(18, 66, 48, 120))
    od.rounded_rectangle((70, 70, 1530, 830), radius=42, outline=(218, 178, 85, 80), width=2)
    for i in range(10):
        x = 1180 + (i * 36)
        od.ellipse((x, 90 + (i % 3) * 10, x + 6, 96 + (i % 3) * 10), fill=(218, 178, 85, 45))

    img = Image.alpha_composite(base, overlay)
    draw = ImageDraw.Draw(img)

    title_font = _load_font(66, bold=True)
    subtitle_font = _load_font(24)
    pill_font = _load_font(16, bold=True)
    small_font = _load_font(18)
    tiny_font = _load_font(14)

    draw.rounded_rectangle((1144, 92, 1468, 132), radius=20, fill=(15, 61, 48, 190), outline=(218, 178, 85, 120), width=2)
    draw.text((1170, 101), "KUMAINA · Propuesta institucional", font=pill_font, fill=(248, 238, 214, 255))

    draw.text((112, 110), "Anexo visual de avance", font=title_font, fill=(244, 232, 209, 255))
    draw.text((114, 178), "Hoja de ruta institucional del MVP cultural y educativo", font=subtitle_font, fill=(225, 195, 117, 255))
    draw.text((114, 214), "De la base pedagógica al piloto: diccionario digital, traductor comunitario y validación territorial.", font=small_font, fill=(229, 226, 216, 255))

    card_left = (112, 276, 496, 694)
    card_mid = (562, 222, 1038, 748)
    card_right = (1104, 276, 1488, 694)
    accent = (180, 132, 32, 255)
    accent_light = (215, 170, 62, 255)

    _draw_card(
        draw,
        card_left,
        "1. Base pedagógica",
        "Lecciones A1/A2/B1",
        [
            "Audio, escritura y práctica breve desde el celular.",
            "Ruta guiada para saludos y frases de uso cotidiano.",
            "Refuerzo lúdico para consolidar vocabulario básico.",
        ],
        0.35,
        accent,
        "Fase 1",
    )

    _draw_card(
        draw,
        card_mid,
        "2. Diccionario comunitario",
        "Consulta viva y validación de contenido",
        [
            "Palabras, expresiones y equivalencias con contexto cultural.",
            "Ejemplos de uso para conservar sentido y memoria.",
            "Base documental para fortalecer identidad lingüística.",
        ],
        0.65,
        accent_light,
        "Fase 2",
        highlight=("Atrinká", "Apretar o cerrar con fuerza, según el contexto."),
    )

    _draw_card(
        draw,
        card_right,
        "3. Traductor + piloto",
        "Prueba con jóvenes y retroalimentación",
        [
            "IA conversacional para práctica oral y traducción.",
            "Métricas de uso, avance y apropiación por nivel.",
            "Piloto con jóvenes y validación territorial.",
        ],
        0.9,
        accent,
        "Fase 3",
        highlight=("¿Cómo tá?", "Senda bien. ¿Y bo?"),
    )

    # Connection line and nodes
    centers = [
        (card_left[2] + 12, (card_left[1] + card_left[3]) // 2),
        ((card_mid[0] + card_mid[2]) // 2, card_mid[1] - 24),
        (card_right[0] - 12, (card_right[1] + card_right[3]) // 2),
    ]
    draw.line([centers[0], (520, 484), centers[1], (1078, 484), centers[2]], fill=(218, 178, 85, 180), width=6, joint="curve")
    for idx, (cx, cy) in enumerate([(520, 484), (800, 188), (1078, 484)], start=1):
        draw.ellipse((cx - 22, cy - 22, cx + 22, cy + 22), fill=accent, outline=(248, 238, 214, 255), width=3)
        w, h = draw.textbbox((0, 0), str(idx), font=_load_font(20, bold=True))[2:]
        draw.text((cx - 7, cy - 12), str(idx), font=_load_font(20, bold=True), fill=(255, 255, 255, 255))

    # Bottom roadmap pill
    pill_box = (150, 770, 1450, 844)
    draw.rounded_rectangle(pill_box, radius=36, fill=(7, 45, 35, 200), outline=(218, 178, 85, 120), width=2)
    draw.text((188, 795), "Ruta de avance", font=_load_font(22, bold=True), fill=(248, 238, 214, 255))
    steps = [("Idea", 0.18), ("Boceto", 0.40), ("MVP", 0.64), ("Piloto", 0.86)]
    for label, pos in steps:
        x = int(pill_box[0] + (pill_box[2] - pill_box[0]) * pos)
        draw.line((x, 786, x, 832), fill=(218, 178, 85, 180), width=2)
        draw.ellipse((x - 8, 800 - 8, x + 8, 800 + 8), fill=(218, 178, 85, 255))
        text_w = draw.textbbox((0, 0), label, font=tiny_font)[2]
        draw.text((x - text_w // 2, 813), label, font=tiny_font, fill=(234, 229, 218, 255))

    draw.text((1128, 742), "Criterio institucional", font=_load_font(16, bold=True), fill=(248, 238, 214, 255))
    draw.text((1094, 764), "La comunidad validará contenido, términos y prioridades.\nLa pieza resume el camino del producto.", font=tiny_font, fill=(225, 218, 202, 255))

    img = img.convert("RGB")
    img.save(output_path, quality=95)
    return output_path


def add_content(doc):
    add_heading(doc, "1. Resumen Ejecutivo", 1)
    doc.add_paragraph(
        "KUMAINA nace de una idea sencilla y urgente: si una lengua se usa en la vida diaria, sigue viva; "
        "si solo aparece en piezas sueltas, se debilita. Por eso la plataforma propone una experiencia "
        "móvil pensada para jóvenes, con lecciones breves, práctica oral, juegos, contexto cultural y una "
        "ruta clara de aprendizaje."
    )
    doc.add_paragraph(
        "La propuesta combina tecnología y memoria cultural sin perder el tono humano. El centro no es la "
        "aplicación por sí misma, sino la comunidad: jóvenes, docentes, sabedores, familias y actores "
        "culturales que ayudan a validar, cuidar y proyectar el contenido."
    )
    doc.add_paragraph(
        "Además, KUMAINA aspira a convertirse en el primer diccionario digital comunitario y traductor "
        "de la lengua palenquera, un puente vivo entre generaciones para consultar palabras, expresiones "
        "y significados sin perder el contexto cultural que les da sentido."
    )
    add_soft_callout(
        doc,
        "Idea fuerza: no se trata solo de traducir palabras; se trata de dar a los jóvenes una razón "
        "para hablar, escuchar y sentir la lengua palenquera como parte de su presente."
    )

    add_heading(doc, "2. Contexto y Oportunidad", 1)
    add_bullet(doc, "La lengua palenquera necesita un espacio digital cotidiano, no solo académico.")
    add_bullet(doc, "Las juventudes aprenden mejor cuando el contenido vive en el celular y se siente cercano.")
    add_bullet(doc, "Existe una oportunidad real de unir patrimonio, educación e innovación en una sola experiencia.")
    add_bullet(doc, "Un diccionario digital y un traductor comunitario pueden ayudar a documentar y compartir la lengua con mayor alcance.")
    add_bullet(doc, "El proyecto puede convertirse en un referente de preservación lingüística con enfoque juvenil.")

    add_heading(doc, "3. Población Objetivo", 1)
    add_bullet(doc, "Jóvenes de 13 a 28 años que quieran aprender o reconectar con la lengua.")
    add_bullet(doc, "Comunidades de San Basilio de Palenque y su diáspora.")
    add_bullet(doc, "Docentes, mediadores, líderes culturales y familias.")
    add_bullet(doc, "Usuarios nacionales e internacionales interesados en lenguas de herencia.")

    add_heading(doc, "4. Objetivo General", 1)
    doc.add_paragraph(
        "Crear una plataforma digital premium para enseñar, preservar y proyectar la lengua palenquera, "
        "incluyendo un diccionario digital comunitario y un traductor cultural, mediante una experiencia "
        "móvil cercana, culturalmente respetuosa y tecnológicamente sólida."
    )

    add_heading(doc, "5. Objetivos Específicos", 1)
    add_bullet(doc, "Diseñar una ruta de aprendizaje por niveles A1, A2 y B1 con progresión clara.")
    add_bullet(doc, "Incorporar práctica de pronunciación, escucha y conversación guiada con apoyo de IA.")
    add_bullet(doc, "Construir el primer diccionario digital comunitario y traductor palenquero, con criterios de validación cultural.")
    add_bullet(doc, "Crear experiencias de juego que refuercen vocabulario, memoria y constancia.")
    add_bullet(doc, "Desarrollar un módulo cultural con historias, música, gastronomía y memoria oral.")
    add_bullet(doc, "Definir métricas de impacto para medir uso, permanencia y aprendizaje real.")

    add_heading(doc, "5.1 Anexo visual de avance", 1)
    doc.add_paragraph(
        "Este anexo sintetiza la ruta de implementación prevista para el producto: una base pedagógica, "
        "el diccionario comunitario, el traductor y el piloto con jóvenes para validar experiencia, "
        "contenido y pertinencia territorial."
    )
    boceto_path = build_boceto_avance_image(Path(__file__).resolve().parent / "kumaina_avance_boceto.png")
    boceto_p = doc.add_paragraph()
    boceto_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    boceto_run = boceto_p.add_run()
    boceto_run.add_picture(str(boceto_path), width=Inches(6.6))
    boceto_caption = doc.add_paragraph()
    boceto_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    boceto_caption_run = boceto_caption.add_run(
        "Anexo visual de la ruta de implementación del proyecto. La imagen resume el avance esperado del MVP."
    )
    boceto_caption_run.italic = True
    boceto_caption_run.font.size = Pt(9)
    boceto_caption_run.font.color.rgb = RGBColor(106, 78, 20)

    add_heading(doc, "6. Alcance Funcional de la Plataforma", 1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["Módulo", "Descripción", "Valor Público"]
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = text
        set_cell_shading(cell, "E8F3EF")
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(13, 67, 52)

    rows = [
        ("Lecciones", "Ruta A1/A2/B1 con audio, escritura y práctica contextual.", "Mejora de competencias comunicativas."),
        ("Diccionario y traductor", "Consulta viva de palabras, expresiones y equivalencias con validación comunitaria.", "Fortalece memoria lingüística y acceso cotidiano."),
        ("IA Conversacional", "Asistente virtual para práctica oral, traducción y corrección.", "Acceso guiado y personalizado."),
        ("Juegos", "Palabra oculta, trivia, ordena la frase y escucha-responde.", "Mayor motivación y permanencia."),
        ("Cultura", "Historias ancestrales, música, gastronomía y referentes.", "Fortalecimiento de identidad cultural."),
        ("Comunidad", "Retos, rankings y participación colaborativa.", "Aprendizaje social y apropiación colectiva."),
    ]
    for row_data in rows:
        row = table.add_row().cells
        for i, value in enumerate(row_data):
            row[i].text = value

    doc.add_paragraph("")

    add_heading(doc, "7. Experiencia de Aprendizaje", 1)
    doc.add_paragraph(
        "La experiencia debe sentirse fluida, clara y útil desde el primer ingreso. Cada lección combina "
        "una explicación corta, una escucha guiada, un ejercicio práctico y una retroalimentación breve. "
        "Eso permite que el usuario avance sin sentirse saturado."
    )
    add_bullet(doc, "A1: palabras base, saludos, familia y frases cotidianas.")
    add_bullet(doc, "A2: conversación contextual, cultura viva y comprensión de expresiones.")
    add_bullet(doc, "B1: fluidez práctica, interpretación y retos guiados.")
    add_bullet(doc, "Cada palabra consultada en el diccionario ayuda a reforzar uso real, memoria y pronunciación.")
    add_bullet(doc, "Cada avance suma XP, racha y reconocimiento visible para motivar continuidad.")

    add_heading(doc, "8. Arquitectura y Tecnología", 1)
    add_bullet(doc, "Frontend móvil: Flutter, pensado para Android e iPhone con una sola base de código.")
    add_bullet(doc, "Backend y autenticación: Firebase, con acceso social, teléfono y control administrativo.")
    add_bullet(doc, "Base de datos: Firestore para progreso, contenidos, usuarios y membresías.")
    add_bullet(doc, "IA y voz: OpenAI API, Speech-to-Text y Text-to-Speech para práctica conversacional.")
    add_bullet(doc, "Distribución: versión web progresiva y publicación posterior en tiendas móviles.")

    add_heading(doc, "9. Cronograma Propuesto (6 meses)", 1)
    cronograma = doc.add_table(rows=1, cols=4)
    cronograma.style = "Table Grid"
    for i, text in enumerate(["Fase", "Duración", "Entregable", "Resultado esperado"]):
        cell = cronograma.rows[0].cells[i]
        cell.text = text
        set_cell_shading(cell, "F4E7C5")
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

    fases = [
        ("Fase 1: Diseño y planeación", "Mes 1", "UX/UI + arquitectura + plan pedagógico", "Base técnica y metodológica aprobada"),
        ("Fase 2: Desarrollo núcleo", "Meses 2-3", "Login, lecciones, progreso y panel admin", "MVP funcional"),
        ("Fase 3: IA + juegos + cultura", "Meses 4-5", "Módulos avanzados y contenidos curados", "Versión beta completa"),
        ("Fase 4: Piloto y ajuste", "Mes 6", "Pruebas con usuarios + optimización", "Versión candidata a despliegue"),
    ]
    for fase in fases:
        row = cronograma.add_row().cells
        for i, value in enumerate(fase):
            row[i].text = value

    doc.add_paragraph("")

    add_heading(doc, "10. Sostenibilidad y Gestión", 1)
    doc.add_paragraph(
        "La sostenibilidad del proyecto depende de una gobernanza sencilla y humana: contenidos validados "
        "por la comunidad, acompañamiento pedagógico, curaduría cultural y ciclos de mejora cortos. La "
        "plataforma debe crecer sin perder su raíz."
    )
    add_bullet(doc, "Comité de contenidos con apoyo de docentes, sabedores y jóvenes.")
    add_bullet(doc, "Actualización del diccionario vivo con nuevas palabras, variantes y ejemplos de uso.")
    add_bullet(doc, "Actualizaciones periódicas de lecciones, audios y retos culturales.")
    add_bullet(doc, "Modelo de acceso por etapas: gratuito, comunitario y premium institucional.")
    add_bullet(doc, "Política de datos y seguridad alineada con buenas prácticas de protección de usuarios.")

    add_heading(doc, "11. Riesgos y Mitigación", 1)
    riesgos = doc.add_table(rows=1, cols=3)
    riesgos.style = "Table Grid"
    for i, text in enumerate(["Riesgo", "Impacto", "Mitigación"]):
        cell = riesgos.rows[0].cells[i]
        cell.text = text
        set_cell_shading(cell, "EAEAEA")
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    riesgo_rows = [
        ("Baja adopción inicial", "Medio", "Piloto con jóvenes y ajustes rápidos de experiencia."),
        ("Contenidos insuficientes", "Alto", "Curaduría comunitaria y producción por lotes."),
        ("Dependencia tecnológica", "Medio", "Arquitectura modular y proveedores reemplazables."),
        ("Uso esporádico", "Medio", "Gamificación, rachas y retos semanales."),
    ]
    for risk in riesgo_rows:
        row = riesgos.add_row().cells
        for i, value in enumerate(risk):
            row[i].text = value

    add_heading(doc, "12. Indicadores de Impacto", 1)
    add_bullet(doc, "Usuarios activos mensuales y retención a 30/90 días.")
    add_bullet(doc, "Lecciones completadas por usuario y tiempo promedio de práctica.")
    add_bullet(doc, "Palabras y expresiones documentadas en el diccionario digital.")
    add_bullet(doc, "Participación en retos, comunidad y módulos culturales.")
    add_bullet(doc, "Percepción de utilidad cultural y educativa en los usuarios piloto.")

    add_heading(doc, "13. Presupuesto Referencial", 1)
    presupuesto = doc.add_table(rows=1, cols=3)
    presupuesto.style = "Table Grid"
    for i, text in enumerate(["Componente", "Descripción", "Valor estimado (COP)"]):
        cell = presupuesto.rows[0].cells[i]
        cell.text = text
        set_cell_shading(cell, "E8F3EF")
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(13, 67, 52)

    rubros = [
        ("Desarrollo tecnológico", "App, backend, integraciones y QA", "$ 180.000.000"),
        ("Producción de contenido", "Lecciones, audios, diccionario digital, traductor y curaduría cultural", "$ 90.000.000"),
        ("IA y analítica", "Modelos, inferencia y monitoreo", "$ 55.000.000"),
        ("Piloto y despliegue", "Pruebas, soporte y operación inicial", "$ 45.000.000"),
        ("Total estimado", "Proyecto integral (6 meses)", "$ 370.000.000"),
    ]
    for rubro in rubros:
        row = presupuesto.add_row().cells
        row[0].text = rubro[0]
        row[1].text = rubro[1]
        row[2].text = rubro[2]

    doc.add_paragraph("")

    add_heading(doc, "14. Solicitud de Articulación con Ministerio TIC", 1)
    doc.add_paragraph(
        "Solicitamos al Ministerio TIC acompañamiento institucional para convertir KUMAINA en una "
        "experiencia piloto de alto impacto que conecte innovación, juventud y patrimonio vivo. "
        "La articulación puede incluir validación técnica, conexión con convocatorias y apoyo en piloto, "
        "especialmente para posicionar el primer diccionario digital comunitario y traductor palenquero."
    )
    add_bullet(doc, "Acompañamiento en estrategia de escalamiento nacional.")
    add_bullet(doc, "Conexión con redes de innovación, educación y juventud.")
    add_bullet(doc, "Visibilidad de buenas prácticas en preservación lingüística digital.")
    add_bullet(doc, "Apoyo para consolidar el diccionario digital y traductor como patrimonio vivo de consulta diaria.")
    add_bullet(doc, "Apoyo en consolidación de sostenibilidad e impacto medible.")

    add_heading(doc, "15. Cierre", 1)
    doc.add_paragraph(
        "KUMAINA quiere dejar una huella clara: que la lengua palenquera se siga escuchando en los "
        "celulares, en la escuela, en la familia y en la comunidad. Este proyecto no busca reemplazar "
        "la voz humana; busca amplificarla con tecnología respetuosa, útil y profundamente conectada "
        "con la identidad de la gente. Y si además dejamos un diccionario digital comunitario y un "
        "traductor vivo, dejamos una herramienta que puede acompañar a la lengua por muchos años más."
    )

    add_heading(doc, "16. Próximos Pasos Propuestos", 1)
    add_bullet(doc, "Mesa técnica inicial con Ministerio TIC para validar alcance y ruta de implementación.")
    add_bullet(doc, "Definición de instituciones piloto (territorio, comunidad y segmento objetivo).")
    add_bullet(doc, "Ajuste de metas trimestrales e indicadores de seguimiento interinstitucional.")
    add_bullet(doc, "Inicio de piloto controlado y reporte de resultados con plan de escalamiento.")

    add_heading(doc, "17. Información de Contacto Institucional", 1)
    contacto = doc.add_table(rows=4, cols=2)
    contacto.style = "Table Grid"
    campos = [
        ("Entidad / Proyecto", "KUMAINA - Lengua Palenquera"),
        ("Representante", "____________________________"),
        ("Correo", "____________________________"),
        ("Teléfono", "____________________________"),
    ]
    for i, (label, value) in enumerate(campos):
        contacto.rows[i].cells[0].text = label
        contacto.rows[i].cells[1].text = value
        set_cell_shading(contacto.rows[i].cells[0], "F4E7C5")
        for p in contacto.rows[i].cells[0].paragraphs:
            for r in p.runs:
                r.bold = True


def main():
    doc = Document()
    configure_document(doc)
    configure_header_footer(doc)
    add_cover(doc)
    add_content(doc)
    output = "docs/Propuesta_KUMAINA_MinTIC.docx"
    doc.save(output)
    print(output)


if __name__ == "__main__":
    main()
